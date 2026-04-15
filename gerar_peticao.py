# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Margens ──────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(3)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2)

# ── Helpers ───────────────────────────────────────────────────────────────────
def add_paragraph(text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  size=12, space_before=0, space_after=6, first_line=True, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.alignment      = align
    p.paragraph_format.space_before   = Pt(space_before)
    p.paragraph_format.space_after    = Pt(space_after)
    if first_line:
        p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    if color:
        run.font.color.rgb = color
    return p

def add_heading(text, level=1, align=WD_ALIGN_PARAGRAPH.CENTER, size=12):
    p = doc.add_paragraph()
    p.paragraph_format.alignment    = align
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    return p

def add_quote(text, size=11):
    """Citação recuada (citação longa)."""
    p = doc.add_paragraph()
    p.paragraph_format.alignment      = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent     = Cm(4)
    p.paragraph_format.right_indent    = Cm(1)
    p.paragraph_format.space_before    = Pt(0)
    p.paragraph_format.space_after     = Pt(6)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    return p

def add_blank(n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run("")
        run.font.size = Pt(6)

# ═══════════════════════════════════════════════════════════════════════════════
# CABEÇALHO
# ═══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(12)
r = p.add_run("EXCELENTÍSSIMO(A) SENHOR(A) DOUTOR(A) JUIZ(A) FEDERAL DO\nJUIZADO ESPECIAL FEDERAL DE HUMAITÁ – AMAZONAS")
r.bold = True
r.font.size = Pt(12)
r.font.name = "Times New Roman"

add_blank()

# Qualificação da parte
add_paragraph(
    "MICHELI PETROSKI FRAGOSO, brasileira, solteira, portadora da Cédula de Identidade "
    "RG nº 32016174 SSP/MT, inscrita no CPF nº 086.877.461-80, residente e domiciliada "
    "na BR 319, KM 170, Zona Rural, Humaitá/AM, CEP 69800-000, por intermédio de sua "
    "advogada infra-assinada, vem, respeitosamente, à presença de Vossa Excelência, "
    "propor a presente:",
    space_after=12
)

add_blank()

# Título da ação
add_heading("AÇÃO PREVIDENCIÁRIA DE CONCESSÃO DE SALÁRIO-MATERNIDADE RURAL", size=13)

add_blank()

add_paragraph(
    "Em face do INSTITUTO NACIONAL DO SEGURO SOCIAL – INSS, autarquia federal, "
    "inscrita no CNPJ sob o nº 29.979.036/0001-40, com sede na Rua Dep. Francisco "
    "Monteiro Neto, Humaitá/AM, CEP 69800-000, pelos seguintes fundamentos fáticos "
    "e jurídicos:",
    space_after=12
)

add_blank()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. JUSTIÇA GRATUITA
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("1. DO PEDIDO DE JUSTIÇA GRATUITA", align=WD_ALIGN_PARAGRAPH.LEFT)

add_paragraph(
    "A Autora é trabalhadora rural em regime de economia familiar, com renda limitada "
    "e voltada à subsistência, declarando, para os devidos fins e sob as penas da lei, "
    "ser pessoa hipossuficiente econômica, sem condições de arcar com as custas "
    "processuais e honorários advocatícios sem prejuízo do próprio sustento e de sua família."
)

add_paragraph(
    "O benefício da gratuidade da justiça encontra amparo no art. 5º, inciso LXXIV, "
    "da Constituição Federal, bem como nos arts. 98 a 102 do Código de Processo Civil "
    "(Lei nº 13.105/2015). A declaração de hipossuficiência, nos termos do art. 99, "
    "§ 3º, do CPC, goza de presunção relativa de veracidade, cabendo à parte contrária "
    "o ônus de afastá-la."
)

add_paragraph(
    "Requer-se, portanto, a concessão dos benefícios da Justiça Gratuita, isentando "
    "a Requerente do pagamento de custas processuais e demais despesas decorrentes "
    "do presente feito."
)

add_blank()

# ═══════════════════════════════════════════════════════════════════════════════
# 2. DOS FATOS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("2. DOS FATOS", align=WD_ALIGN_PARAGRAPH.LEFT)

add_paragraph(
    "A Autora, Sra. Micheli Petroski Fragoso, reside em união estável com seu companheiro, "
    "ROBSON LOURENÇO DE SOUZA, na zona rural do Município de Humaitá/AM. Antes de se "
    "estabelecerem em Humaitá/AM, o casal residia na propriedade dos avós paternos da "
    "criança, em Rondolândia/MT — endereço que, por essa razão, consta na certidão de "
    "nascimento."
)

add_paragraph(
    "Desde agosto de 2022, o casal reside e trabalha na propriedade rural do sogro da "
    "Autora (avô paterno da criança), Sr. PAULO GABRIEL LOURENÇO, onde todos se dedicam "
    "à atividade agropecuária em regime de economia familiar, nos termos do art. 11, "
    "inciso VII, alínea 'c', da Lei nº 8.213/91."
)

add_paragraph(
    "A Autora engravidou em junho de 2023, e o nascimento de sua filha, MARIA GABRIELI "
    "LOURENÇO FRAGOSO, ocorreu em 13 de março de 2024. Em 14/02/2025, a Autora "
    "protocolou requerimento administrativo de salário-maternidade rural perante o INSS "
    "(Protocolo nº 1186879712). Contudo, a autarquia previdenciária indeferiu o pedido "
    "sob a alegação de 'ausência de comprovação da qualidade de segurada especial na "
    "data do afastamento'."
)

add_paragraph(
    "Tal decisão administrativa é manifestamente equivocada e desconsidera a realidade "
    "fática e documental apresentada. A Autora sempre residiu em área rural, em terras "
    "integrantes do núcleo familiar de seu companheiro, onde exerce atividade pecuária "
    "e agrícola indispensável à subsistência da família."
)

add_paragraph(
    "O fato de o parto ter ocorrido em Cacoal/RO não descaracteriza a condição de "
    "segurada especial da Autora, uma vez que o deslocamento se deu exclusivamente em "
    "razão da busca por melhor infraestrutura hospitalar, mantendo-se intacto o vínculo "
    "com a atividade rural desempenhada em Humaitá/AM. Tal situação é corriqueira no "
    "interior do País e não pode constituir fundamento para indeferimento do benefício."
)

add_paragraph(
    "Importa destacar que a residência da Autora não conta com energia elétrica "
    "convencional, valendo-se apenas de painel solar — circunstância absolutamente "
    "comum em localidades rurais remotas. Por essa razão, os comprovantes de endereço "
    "estão em nome de seu sogro, o que, conforme pacificado na jurisprudência pátria, "
    "não constitui óbice à comprovação da residência e da atividade rurícola."
)

add_paragraph(
    "O acompanhamento pré-natal da Autora foi realizado integralmente em Humaitá/AM, "
    "conforme atesta a Ficha de Cadastro do Usuário da Secretaria Municipal de Saúde "
    "de Humaitá/AM, datada de 18/08/2023, onde constam o endereço rural (BR 319, "
    "KM 170) e a anotação de 'gestante', documento de especial relevância para a "
    "comprovação da condição de segurada especial no período de carência."
)

add_blank()

add_heading("2.1. Dados do Requerimento Administrativo", align=WD_ALIGN_PARAGRAPH.LEFT)

add_paragraph(
    "• Protocolo do Requerimento Administrativo: nº 1186879712\n"
    "• Data do Protocolo: 14/02/2025\n"
    "• Data do Parto: 13/03/2024\n"
    "• Número do Benefício (NB): 209.147.265-9\n"
    "• Decisão Administrativa: Indeferimento – 'ausência de comprovação da qualidade "
    "de segurada especial na data do afastamento'",
    first_line=False
)

add_blank()

# ═══════════════════════════════════════════════════════════════════════════════
# 3. DO DIREITO
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("3. DO DIREITO E DAS PROVAS", align=WD_ALIGN_PARAGRAPH.LEFT)

add_heading("3.1. Do Fundamento Constitucional e Legal", align=WD_ALIGN_PARAGRAPH.LEFT)

add_paragraph(
    "O direito reclamado pela Requerente encontra amparo na Constituição Federal, na "
    "legislação previdenciária e na jurisprudência pátria."
)

add_paragraph(
    "A Constituição Federal de 1988 assegura, em seu art. 201, inciso II, a proteção "
    "à maternidade como um dos pilares da Previdência Social:"
)

add_quote(
    '"Art. 201. A previdência social será organizada sob a forma do Regime Geral de '
    'Previdência Social, de caráter contributivo e de filiação obrigatória, observados '
    'critérios que preservem o equilíbrio financeiro e atuarial, e atenderá, na forma '
    'da lei, a: [...] II – proteção à maternidade, especialmente à gestante."'
)

add_paragraph(
    "Igualmente, o art. 7º, inciso XVIII, da Constituição Federal garante:"
)

add_quote(
    '"Art. 7º São direitos dos trabalhadores urbanos e rurais, além de outros que '
    'visem à melhoria de sua condição social: [...] XVIII – licença à gestante, sem '
    'prejuízo do emprego e do salário, com a duração de cento e vinte dias."'
)

add_paragraph(
    "No plano infraconstitucional, a Lei nº 8.213/91 define a segurada especial e os "
    "requisitos para a concessão do salário-maternidade rural:"
)

add_quote(
    '"Art. 11. São segurados obrigatórios da previdência social as seguintes pessoas '
    'físicas: [...] VII – como segurado especial: a pessoa física residente no imóvel '
    'rural ou em aglomerado urbano ou rural próximo a ele que, individualmente ou em '
    'regime de economia familiar, ainda que com o auxílio eventual de terceiro, na '
    'condição de: a) produtor, seja proprietário, usufrutuário, possuidor, assentado, '
    'parceiro ou meeiro outorgados, comodatário ou arrendatário rurais, que explore '
    'atividade agropecuária em área de até 4 (quatro) módulos fiscais; [...] '
    'c) cônjuge ou companheiro, bem como filho maior de 16 (dezesseis) anos de idade '
    'ou a este equiparado, do segurado de que tratam as alíneas a e b deste inciso, '
    'que, comprovadamente, trabalhem com o grupo familiar respectivo."'
)

add_paragraph(
    "Quanto ao salário-maternidade da segurada especial, o art. 39, parágrafo único, "
    "combinado com o art. 71-A, § 2º, da Lei nº 8.213/91, e o art. 93, § 2º, do "
    "Decreto nº 3.048/99 (Regulamento da Previdência Social), estabelecem:"
)

add_quote(
    '"§ 2º Será devido o salário-maternidade à segurada especial, desde que comprove '
    'o exercício de atividade rural nos últimos 10 (dez) meses imediatamente anteriores '
    'à data do parto ou do requerimento do benefício, quando requerido antes do parto, '
    'mesmo que de forma descontínua."'
)

add_paragraph(
    "Outrossim, registra-se que o direito ao benefício não está sujeito a prazo "
    "decadencial de 180 dias, submetendo-se apenas ao prazo prescricional de 5 (cinco) "
    "anos, conforme entendimento pacificado pelo Superior Tribunal de Justiça e pela "
    "Turma Nacional de Uniformização dos Juizados Especiais Federais (TNU)."
)

add_blank()

add_heading("3.2. Do Início de Prova Material em Nome do Grupo Familiar", align=WD_ALIGN_PARAGRAPH.LEFT)

add_paragraph(
    "É pacífico o entendimento de que documentos em nome de outros membros do grupo "
    "familiar (companheiro, pai ou sogro) servem como início de prova material para a "
    "trabalhadora rural. Nesse sentido, a Súmula 73 do egrégio Tribunal Regional "
    "Federal da 4ª Região estabelece:"
)

add_quote(
    '"Admitem-se como início de prova material do efetivo exercício de atividade rural, '
    'em regime de economia familiar, documentos de terceiros, membros do grupo familiar, '
    'incluídos filhos e irmãos."'
)

add_paragraph(
    "No mesmo sentido, a jurisprudência consolidada do Superior Tribunal de Justiça, "
    "em sede de recurso repetitivo (REsp 1.304.479/SP, Rel. Min. Herman Benjamin, "
    "1ª Seção, j. 22/05/2013), reafirma que o trabalho da mulher no campo é, via de "
    "regra, exercido em regime de colaboração mútua, e que documentos em nome de "
    "membros do grupo familiar configuram início razoável de prova material, a ser "
    "corroborado pela prova testemunhal."
)

add_paragraph(
    "Para comprovar a qualidade de segurada especial, a Autora apresenta os seguintes "
    "documentos:"
)

add_paragraph(
    "a) Contrato Particular de Compra e Venda (21/10/2022): Comprova a aquisição do "
    "\"Sítio Santo Antônio\", Gleba Acará, em Humaitá/AM, pelo sogro da Autora, Sr. "
    "Paulo Gabriel Lourenço, com firma reconhecida em cartório, demonstrando a posse "
    "da terra pelo núcleo familiar muito antes do período de carência.",
    first_line=False
)

add_paragraph(
    "b) Guia de Trânsito Animal – GTA (22/01/2024): Emitida pela ADAF/Amazonas, em "
    "nome do sogro da Autora, comprovando a movimentação de bovinos (fêmeas adultas) "
    "para fins de recria. Este documento insere-se exatamente dentro do período de "
    "carência (10 meses anteriores ao parto de março/2024).",
    first_line=False
)

add_paragraph(
    "c) Comprovante de Aquisição de Ferramenta Rural (22/05/2024): Nota de compra de "
    "motosserra profissional em nome do companheiro da Autora (Robson Lourenço de "
    "Souza), reforçando o vínculo com a localidade rural (BR 319, KM 170 – Realidade).",
    first_line=False
)

add_paragraph(
    "d) Certidão de Nascimento da Filha (MARIA GABRIELI LOURENÇO FRAGOSO): Constando "
    "a residência dos pais em zona rural.",
    first_line=False
)

add_paragraph(
    "e) Ficha de Cadastro do Usuário (18/08/2023): Emitida pela Secretaria Municipal "
    "de Saúde de Humaitá/AM, comprovando o acompanhamento pré-natal da Autora na "
    "localidade rural, com registro do endereço (BR 319, KM 170) e a anotação de "
    "\"gestante\", em período crucial para a comprovação da qualidade de segurada especial.",
    first_line=False
)

add_blank()

add_heading("3.3. Da Carência e da Qualidade de Segurada Especial", align=WD_ALIGN_PARAGRAPH.LEFT)

add_paragraph(
    "Considerando que o parto ocorreu em 13/03/2024, o período de carência de dez "
    "meses compreende o intervalo entre maio de 2023 e março de 2024. O Contrato de "
    "Compra e Venda de outubro/2022 e a GTA de janeiro/2024 cobrem integralmente esse "
    "lapso temporal, demonstrando que a Autora, inserida no regime de economia familiar, "
    "preenche todos os requisitos legais para a concessão do benefício."
)

add_paragraph(
    "Ressalta-se, ainda, que a Ficha de Cadastro do SUS, datada de 18/08/2023, "
    "situa-se precisamente dentro do período de carência, constituindo prova robusta "
    "da permanência da Autora na zona rural e de sua condição de segurada especial "
    "durante toda a gestação."
)

add_blank()

add_heading("3.4. Da Jurisprudência Aplicável", align=WD_ALIGN_PARAGRAPH.LEFT)

add_paragraph(
    "A jurisprudência pátria é uníssona no reconhecimento do direito ao salário-"
    "maternidade em casos análogos ao presente. Confira-se:"
)

add_quote(
    '"PREVIDENCIÁRIO. SALÁRIO-MATERNIDADE. TRABALHADORA RURAL. SEGURADA ESPECIAL. '
    'ATIVIDADE CAMPESINA COMPROVADA. INÍCIO DE PROVA MATERIAL. PROVA TESTEMUNHAL. '
    'CONCESSÃO DO BENEFÍCIO. SENTENÇA DE IMPROCEDÊNCIA REFORMADA. 1. O salário-'
    'maternidade é devido à segurada especial, no valor de 1 (um) salário mínimo '
    'mensal durante 120 dias, a contar da data do parto ou dos 28 (vinte e oito) dias '
    'que o antecederam, desde que comprovado o exercício de atividade rural, ainda que '
    'de forma descontínua, nos dez meses imediatamente anteriores ao início do benefício '
    '(arts. 39, parágrafo único, e 71 c/c art. 25 da Lei nº 8.213/91). 2. Na hipótese, '
    'a parte autora comprovou sua condição de segurada especial durante o período de '
    'carência mediante documentos em nome de membros do grupo familiar, que, analisados '
    'em conjunto com a prova testemunhal, configuram início razoável de prova material '
    'suficiente à concessão do benefício. 3. Recurso provido."'
)

add_paragraph(
    "No mesmo sentido, é o entendimento da Turma Nacional de Uniformização dos "
    "Juizados Especiais Federais (TNU), que, por meio da Súmula nº 41, consolidou "
    "o entendimento de que:"
)

add_quote(
    'A circunstância de um dos integrantes do núcleo familiar desempenhar atividade '
    'urbana não implica, por si só, a descaracterização do regime de economia familiar.'
)

add_blank()

# ═══════════════════════════════════════════════════════════════════════════════
# 4. DOS PEDIDOS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("4. DOS PEDIDOS", align=WD_ALIGN_PARAGRAPH.LEFT)

add_paragraph("Diante de todo o exposto, requer-se a Vossa Excelência:")

add_paragraph(
    "a) A concessão dos benefícios da Justiça Gratuita, nos termos do art. 99, "
    "§ 3º, do CPC, diante da hipossuficiência econômica da Autora;",
    first_line=False
)

add_paragraph(
    "b) A citação do INSS, na pessoa de seu representante legal, para, querendo, "
    "apresentar contestação no prazo legal;",
    first_line=False
)

add_paragraph(
    "c) A produção de todas as provas admitidas em Direito, especialmente a prova "
    "testemunhal, essencial para corroborar o início de prova material apresentado "
    "e para detalhar a rotina de labor rural da Autora;",
    first_line=False
)

add_paragraph(
    "d) O julgamento de PROCEDÊNCIA TOTAL da demanda, condenando o INSS a:",
    first_line=False
)

add_paragraph(
    "    d.1) Conceder à Autora o benefício de Salário-Maternidade Rural "
    "(NB 209.147.265-9), com implantação imediata;",
    first_line=False
)

add_paragraph(
    "    d.2) Efetuar o pagamento das parcelas vencidas desde a data do parto "
    "(13/03/2024), devidamente corrigidas pelo INPC e acrescidas de juros de mora "
    "nos termos do art. 1º-F da Lei nº 9.494/97, com a redação dada pela Lei "
    "nº 11.960/2009;",
    first_line=False
)

add_paragraph(
    "e) A condenação do INSS ao pagamento de honorários advocatícios, na hipótese "
    "de interposição de recurso, nos termos do art. 55 da Lei nº 9.099/95, combinado "
    "com o art. 1º da Lei nº 10.259/2001.",
    first_line=False
)

add_blank()

# Valor da causa
add_paragraph(
    "Dá-se à causa o valor de R$ 7.248,00 (sete mil, duzentos e quarenta e oito reais), "
    "correspondente a 6 (seis) parcelas de salário-maternidade no valor do salário "
    "mínimo vigente (R$ 1.518,00 em 2025), calculadas a partir da data do parto "
    "(13/03/2024), para fins de alçada do Juizado Especial Federal, nos termos "
    "do art. 3º da Lei nº 10.259/2001.",
    bold=False
)

add_blank()

add_paragraph(
    "Termos em que, pede deferimento.",
    align=WD_ALIGN_PARAGRAPH.CENTER,
    first_line=False
)

add_blank()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Humaitá/AM, 06 de abril de 2026.")
r.font.size = Pt(12)
r.font.name = "Times New Roman"

add_blank(2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("_________________________________________________\n"
              "HEVELLYN PRYSCYLLA MEDEIROS ROBERTO\n"
              "Advogada – OAB/RO nº 6.595")
r.font.size = Pt(12)
r.font.name = "Times New Roman"
r.bold = True

# ── Salvar ────────────────────────────────────────────────────────────────────
output = "/vercel/sandbox/uploads/petição - michele salario maternidade - CORRIGIDA.docx"
doc.save(output)
print(f"Documento salvo em: {output}")
